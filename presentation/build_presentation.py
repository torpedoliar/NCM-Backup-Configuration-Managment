import os
from pathlib import Path

# Try imports and provide helpful error if missing
try:
    from PIL import Image, ImageDraw, ImageFont
except Exception as e:
    raise SystemExit("Pillow (PIL) is required. Install with: pip install pillow")

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception as e:
    raise SystemExit("python-docx is required. Install with: pip install python-docx")

ROOT = Path(__file__).resolve().parent.parent
PRES_DIR = ROOT / "presentation"
IMG_DIR = PRES_DIR / "images"
IMG_DIR.mkdir(parents=True, exist_ok=True)

# Simple drawing helpers ------------------------------------------------------

def load_font(size: int):
    try:
        return ImageFont.truetype("arial.ttf", size)
    except Exception:
        return ImageFont.load_default()

def text_center(draw: ImageDraw.ImageDraw, xy, text, font, fill=(0,0,0)):
    x1, y1, x2, y2 = xy
    w = x2 - x1
    h = y2 - y1
    tw, th = draw.textbbox((0,0), text, font=font)[2:]
    tx = x1 + (w - tw) // 2
    ty = y1 + (h - th) // 2
    draw.text((tx, ty), text, font=font, fill=fill)

def draw_box(draw, xy, text, fill=(255,255,255), outline=(0,0,0)):
    draw.rounded_rectangle(xy, radius=8, outline=outline, width=2, fill=fill)
    text_center(draw, xy, text, load_font(16))

def draw_title(draw, text, canvas_w):
    f = load_font(22)
    tw, th = draw.textbbox((0,0), text, font=f)[2:]
    draw.text(((canvas_w - tw)//2, 10), text, font=f, fill=(10,10,10))

def draw_arrow(draw, start, end, color=(0,0,0)):
    # line
    draw.line([start, end], fill=color, width=3)
    # simple arrow head
    import math
    ang = math.atan2(end[1]-start[1], end[0]-start[0])
    sz = 10
    left = (end[0] - sz*math.cos(ang - math.pi/6), end[1] - sz*math.sin(ang - math.pi/6))
    right = (end[0] - sz*math.cos(ang + math.pi/6), end[1] - sz*math.sin(ang + math.pi/6))
    draw.polygon([end, left, right], fill=color)

# Diagram generators ----------------------------------------------------------

def gen_backup_process_png(path: Path):
    W, H = 1100, 800
    im = Image.new("RGB", (W, H), (248, 250, 252))
    d = ImageDraw.Draw(im)
    draw_title(d, "Backup Process Flow", W)

    # positions (x1,y1,x2,y2)
    boxes = {
        "A": (60, 80, 260, 140),
        "B": (320, 80, 540, 160),
        "C": (620, 60, 860, 120),
        "D": (620, 140, 860, 200),
        "E": (60, 220, 260, 280),
        "F": (320, 220, 540, 280),
        "G": (620, 220, 860, 280),
        "H": (60, 320, 260, 380),
        "I": (320, 320, 540, 380),
        "J": (620, 320, 860, 380),
        "K": (60, 420, 260, 480),
        "L": (320, 420, 540, 480),
        "M": (620, 420, 860, 480),
        "N": (60, 520, 260, 580),
        "O": (320, 520, 540, 580),
    }
    labels = {
        "A": "Trigger Backup",
        "B": "Manual / Scheduled?",
        "C": "GUI Action",
        "D": "Background Service",
        "E": "Load Switch + Credentials",
        "F": "Decrypt Credentials",
        "G": "Select Protocol",
        "H": "SSHClient Connect",
        "I": "TelnetClient Connect",
        "J": "Execute Commands",
        "K": "Capture Output",
        "L": "Save File",
        "M": "SHA256 Hash + DB Record",
        "N": "Compare with Last",
        "O": "Generate Diff / Identical",
    }

    # draw boxes
    for k, rect in boxes.items():
        draw_box(d, rect, labels[k])

    # arrows (rough path)
    def center(rect):
        x1,y1,x2,y2 = rect
        return ((x1+x2)//2, (y1+y2)//2)

    draw_arrow(d, center(boxes["A"]), center(boxes["B"]))
    draw_arrow(d, (boxes["B"][2], center(boxes["B"])[1]-20), (boxes["C"][0], center(boxes["C"])[1]))
    draw_arrow(d, (boxes["B"][2], center(boxes["B"])[1]+20), (boxes["D"][0], center(boxes["D"])[1]))
    draw_arrow(d, center(boxes["C"]), center(boxes["E"]))
    draw_arrow(d, center(boxes["D"]), center(boxes["E"]))
    draw_arrow(d, center(boxes["E"]), center(boxes["F"]))
    draw_arrow(d, center(boxes["F"]), center(boxes["G"]))
    draw_arrow(d, (boxes["G"][0], boxes["G"][3]+10), (boxes["H"][2], boxes["H"][1]))
    draw_arrow(d, center(boxes["G"]), center(boxes["I"]))
    draw_arrow(d, center(boxes["H"]), center(boxes["J"]))
    draw_arrow(d, center(boxes["I"]), center(boxes["J"]))
    draw_arrow(d, center(boxes["J"]), center(boxes["K"]))
    draw_arrow(d, center(boxes["K"]), center(boxes["L"]))
    draw_arrow(d, center(boxes["L"]), center(boxes["M"]))
    draw_arrow(d, center(boxes["M"]), center(boxes["N"]))
    draw_arrow(d, center(boxes["N"]), center(boxes["O"]))

    im.save(path)


def gen_architecture_png(path: Path):
    W, H = 1000, 700
    im = Image.new("RGB", (W, H), (248, 250, 252))
    d = ImageDraw.Draw(im)
    draw_title(d, "System Architecture (Layered)", W)

    layer_w, layer_h = 860, 100
    x = (W - layer_w)//2
    gaps = 40
    y0 = 90

    layers = [
        ("UI Layer (ttkbootstrap): Dashboard, Inventory, History, Settings", (x, y0, x+layer_w, y0+layer_h)),
        ("Service Layer: BackupService, ScheduleService, CryptoService, DiffService", (x, y0+(layer_h+gaps), x+layer_w, y0+(layer_h+gaps)+layer_h)),
        ("Network Layer: SSHClient, TelnetClient, BackupRunner", (x, y0+2*(layer_h+gaps), x+layer_w, y0+2*(layer_h+gaps)+layer_h)),
        ("Data Layer (SQLite + SQLAlchemy): Credentials, Devices, Backups", (x, y0+3*(layer_h+gaps), x+layer_w, y0+3*(layer_h+gaps)+layer_h)),
    ]
    for text, rect in layers:
        draw_box(d, rect, text, fill=(235, 245, 255))

    # arrows down
    for i in range(len(layers)-1):
        r1 = layers[i][1]
        r2 = layers[i+1][1]
        draw_arrow(d, ((r1[0]+r1[2])//2, r1[3]), ((r2[0]+r2[2])//2, r2[1]))

    im.save(path)


def gen_startup_sequence_png(path: Path):
    W, H = 1100, 600
    im = Image.new("RGB", (W, H), (248, 250, 252))
    d = ImageDraw.Draw(im)
    draw_title(d, "Startup Sequence (Simplified)", W)

    # lifelines
    cols = [120, 380, 640, 900]
    actors = ["User", "GUI", "Service", "Scheduler"]
    for i, x in enumerate(cols):
        d.text((x-20, 60), actors[i], font=load_font(16), fill=(0,0,0))
        d.line([(x, 90), (x, H-40)], fill=(120,120,120), width=2)

    def msg(fr, to, y, text):
        x1 = cols[fr]
        x2 = cols[to]
        draw_arrow(d, (x1, y), (x2, y))
        d.text(((x1+x2)//2 - 100, y-18), text, font=load_font(14), fill=(0,0,0))

    y = 130
    msg(0,1,y,"Launch EXE"); y+=40
    d.text((cols[1]-150, y), "[First Run] Setup passphrase / [Normal] Unlock vault", font=load_font(13), fill=(0,0,0)); y+=40
    msg(1,3,y,"Start scheduler"); y+=40
    msg(3,2,y,"Execute scheduled backups"); y+=40
    d.text((cols[1]-140, y), "GUI updates dashboard every 30s", font=load_font(13), fill=(0,0,0)); y+=60
    msg(0,1,y,"Close window"); y+=40
    d.text((cols[1]-160, y), "Stop scheduler (if not service mode)", font=load_font(13), fill=(0,0,0))

    im.save(path)


def gen_windows_service_png(path: Path):
    W, H = 1100, 600
    im = Image.new("RGB", (W, H), (248, 250, 252))
    d = ImageDraw.Draw(im)
    draw_title(d, "Windows Service Integration", W)

    boxes = {
        "boot": (80, 110, 300, 170),
        "ts": (380, 110, 650, 170),
        "exe": (760, 110, 1020, 170),
        "lockq": (380, 230, 650, 290),
        "lock": (760, 230, 1020, 290),
        "sched": (380, 350, 650, 410),
        "jobs": (760, 350, 1020, 410),
    }
    labels = {
        "boot": "Windows Boot",
        "ts": "Task Scheduler (SYSTEM)",
        "exe": "Service EXE --service",
        "lockq": "Lock exists?",
        "lock": "Acquire/Validate Lock",
        "sched": "Start Scheduler",
        "jobs": "Run Scheduled Jobs",
    }
    for k, rect in boxes.items():
        draw_box(d, rect, labels[k], fill=(240, 255, 240) if k in ("exe","jobs") else (255,255,255))

    def c(rect):
        x1,y1,x2,y2 = rect
        return ((x1+x2)//2, (y1+y2)//2)

    draw_arrow(d, c(boxes["boot"]), c(boxes["ts"]))
    draw_arrow(d, c(boxes["ts"]), c(boxes["exe"]))
    draw_arrow(d, c(boxes["exe"]), c(boxes["lockq"]))
    draw_arrow(d, c(boxes["lockq"]), c(boxes["lock"]))
    draw_arrow(d, c(boxes["lock"]), c(boxes["sched"]))
    draw_arrow(d, c(boxes["sched"]), c(boxes["jobs"]))

    im.save(path)

# Build all images ------------------------------------------------------------
backup_png = IMG_DIR / "backup_process.png"
arch_png = IMG_DIR / "system_architecture.png"
startup_png = IMG_DIR / "startup_sequence.png"
service_png = IMG_DIR / "windows_service_integration.png"

gen_backup_process_png(backup_png)
gen_architecture_png(arch_png)
gen_startup_sequence_png(startup_png)
gen_windows_service_png(service_png)

# Build DOCX ------------------------------------------------------------------

doc = Document()

doc.add_heading('Allied Telesis Backup Configuration Manager', 0)
sub = doc.add_paragraph('Presentasi Teknis v3.5.4 (Production)')
sub.runs[0].bold = True

# Latar Belakang
doc.add_heading('Latar Belakang', level=1)
doc.add_paragraph(
    "Di lingkungan SJA, jumlah switch yang digunakan cukup banyak dan setiap switch memiliki konfigurasi unik. "
    "Tidak ada mekanisme backup otomatis dan terpusat (selama ini manual), sehingga menjadi pain point ketika terjadi "
    "kehilangan konfigurasi atau perlu penggantian perangkat. Dibutuhkan solusi yang sederhana, andal, dan mudah dikelola."
)

# Solusi
doc.add_heading('Solusi yang Dipertimbangkan', level=1)
doc.add_paragraph('1) Menggunakan Allied Telesis Vista Manager: manajemen terpusat, discovery/topologi, monitoring, backup terintegrasi, dan fitur enterprise lainnya. (Contoh tampilan tidak disertakan).')
doc.add_paragraph('2) Menggunakan aplikasi internal "ATI Backup Configuration Management" (project ini): berfokus pada fitur yang krusial untuk operasional SJA – backup terjadwal, penyimpanan terstruktur, diff konfigurasi, enkripsi kredensial, dan histori. Fitur yang tidak esensial untuk kebutuhan saat ini dikecilkan agar solusi tetap ringan dan portable.')

doc.add_paragraph('Perbedaan: Vista Manager menyediakan fitur enterprise yang lebih luas; aplikasi internal memfokuskan pada kebutuhan inti (operasional harian) untuk efisiensi, kemudahan deployment, dan kontrol data on-premises.')

# Vista Manager details
doc.add_heading('Allied Telesis Vista Manager (ATI) - Fitur & Contoh Tampilan', level=2)
for line in [
    'Fitur: Discovery dan topologi jaringan otomatis (map interaktif).',
    'Fitur: Inventori perangkat, status kesehatan, dan firmware management.',
    'Fitur: Backup/restore konfigurasi terintegrasi dengan kebijakan global.',
    'Fitur: Dashboard monitoring, alarm & notifikasi terpusat.',
    'Fitur: Otentikasi terpusat dan role-based access.',
]:
    doc.add_paragraph(line, style='List Bullet')

doc.add_paragraph('Contoh tampilan (deskriptif): Topology map interaktif dengan node switch, panel inventori perangkat (grid) dengan filter/status, halaman kebijakan backup (jadwal & retensi), serta dashboard ringkas berisi grafik alarm dan job status.')

doc.add_heading('Perbandingan Singkat', level=2)
for line in [
    'Vista Manager: fitur enterprise yang luas (monitoring full, topology, multi-site).',
    'Aplikasi internal: fokus pada operasi inti (backup terjadwal, diff, retensi, enkripsi), ringan & portable.',
    'Vista Manager: membutuhkan infrastruktur dan lisensi; aplikasi internal: distribusi EXE, mudah dipindah & on-prem.',
]:
    doc.add_paragraph(line, style='List Bullet')

doc.add_heading('Kelebihan Aplikasi "ATI Backup Configuration Management"', level=2)
advantages = [
    'Fokus pada kebutuhan SJA: hanya fitur krusial (backup terjadwal, diff, retensi, enkripsi) → sederhana & cepat dipakai.',
    'Portable & zero-install: EXE tunggal, base_dir dinamis mengikuti lokasi EXE; mudah dipindah tanpa re-setup.',
    'Ringan (low footprint): SQLite embedded, tidak butuh DB/server terpisah.',
    'On-prem & offline: data kredensial/backup tersimpan lokal, tetap berjalan tanpa internet.',
    'Efisien biaya: tanpa lisensi vendor; mudah dirawat oleh tim internal.',
    'Multi-protocol: SSH & Telnet untuk perangkat modern dan legacy.',
    'Keamanan kredensial: Fernet/AES-128 + PBKDF2-HMAC-SHA256; hash SHA-256 untuk verifikasi integritas backup.',
    'Scheduling handal: tetap berjalan di background (Task Scheduler/Windows Service) meski GUI ditutup.',
    'Retention otomatis: kebijakan 365 hari menjaga storage tetap sehat.',
    'Diff & audit trail: deteksi perubahan cepat dengan histori dan logging terstruktur.',
    'Fleksibel dikembangkan: perubahan/fitur baru dapat ditambahkan cepat sesuai kebutuhan SJA.',
]
for a in advantages:
    doc.add_paragraph(a, style='List Bullet')

# Fitur
doc.add_heading('Fitur Utama', level=1)
features = [
    'Multi-Protocol: SSH & Telnet',
    'Scheduling Otomatis: 15m, 1h, 6h, 12h, 24h',
    'Enkripsi kredensial (Fernet/AES-128 + PBKDF2-HMAC-SHA256)',
    'Configuration Diff & Riwayat',
    'Retention Policy 365 hari',
    'Windows Service / Task Scheduler',
    'Portable EXE (base_dir dinamis mengikuti lokasi EXE)',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

# Struktur Project
doc.add_heading('Struktur Project (Ringkas)', level=1)
code = (
"app/\n"
"  main.py                # Entry point GUI/service\n"
"  windows_service.py     # Wrapper service\n"
"  config/                # Konfigurasi & logging\n"
"  net/                   # SSH/Telnet clients\n"
"  services/              # Backup/Crypto/Diff/Retention\n"
"  ui/                    # Dashboard/Inventory/History/Settings\n"
"backups/  data/  logs/   # Output & runtime data\n"
)
p = doc.add_paragraph()
r = p.add_run(code)
r.font.name = 'Consolas'
r.font.size = Pt(10)

# Flowcharts (Images)
doc.add_heading('Flowchart - Backup Process', level=1)
doc.add_picture(str(backup_png), width=Inches(6.5))

doc.add_heading('Flowchart - System Architecture', level=1)
doc.add_picture(str(arch_png), width=Inches(6.5))

doc.add_heading('Flowchart - Startup Sequence', level=1)
doc.add_picture(str(startup_png), width=Inches(6.5))

doc.add_heading('Flowchart - Windows Service Integration', level=1)
doc.add_picture(str(service_png), width=Inches(6.5))

# Teknologi
doc.add_heading('Teknologi Utama', level=1)
stack = [
    'Python 3, ttkbootstrap (UI)',
    'Paramiko (SSH) & telnetlib3 (Telnet)',
    'SQLAlchemy + SQLite',
    'APScheduler (penjadwalan background)',
    'cryptography (Fernet/AES-128, PBKDF2-HMAC-SHA256)',
    'PyInstaller (distribusi EXE)',
]
for s in stack:
    doc.add_paragraph(s, style='List Bullet')

# Keamanan & Deployment
doc.add_heading('Keamanan & Deployment', level=1)
doc.add_paragraph('• Kredensial terenkripsi (Fernet/AES-128) dengan kunci yang diturunkan dari passphrase melalui PBKDF2-HMAC-SHA256.')
doc.add_paragraph('• Hash SHA-256 untuk verifikasi integritas file konfigurasi.')
doc.add_paragraph('• Build via PyInstaller; aplikasi portable; integrasi Task Scheduler/Windows Service.')

# Save
out_path = ROOT / 'presentasi_v2.docx'
doc.save(str(out_path))
print(f"Presentasi berhasil dibuat: {out_path}")
