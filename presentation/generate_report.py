from docx import Document
from docx.shared import Inches
import os

doc = Document()

# Add title
doc.add_heading('Allied Telesis Backup Configuration Manager', 0)

# Add subtitle
doc.add_heading('Technical Report v3.5.4', level=1)

# Add summary
doc.add_heading('Executive Summary', level=2)
doc.add_paragraph(
    "Aplikasi desktop Windows untuk backup otomatis konfigurasi switch Allied Telesis dengan fitur:\n"
    "- Multi-protocol (SSH/Telnet)\n"
    "- Enkripsi AES-128 untuk kredensial\n"
    "- Scheduling fleksibel\n"
    "- Portable deployment\n"
    "- Windows Service integration"
)

# Add features
doc.add_heading('Key Features', level=2)
doc.add_picture('feature_pie_chart.png', width=Inches(6))
doc.add_paragraph("Diagram distribusi fitur utama aplikasi")

# Add more sections similarly...

# Save document
doc.save('report.docx')
