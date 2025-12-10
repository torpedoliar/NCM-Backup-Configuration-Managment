from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

# Create new document
doc = Document()

# Add title
doc.add_heading('Allied Telesis Backup Configuration Manager', 0)

# Add subtitle
doc.add_heading('Technical Report v3.5.4', level=1)

# Add summary section
doc.add_heading('Executive Summary', level=2)
doc.add_paragraph(
    "Aplikasi desktop Windows untuk backup otomatis konfigurasi switch Allied Telesis dengan fitur:"
    "\n- Multi-protocol (SSH/Telnet)"
    "\n- Enkripsi AES-128 untuk kredensial"
    "\n- Scheduling fleksibel"
    "\n- Portable deployment"
    "\n- Windows Service integration"
)

# Add features section
doc.add_heading('Key Features', level=2)
doc.add_paragraph("• Multi-Protocol Support: SSH dan Telnet")
doc.add_paragraph("• Scheduling Otomatis: Interval 15m, 1h, 6h, 12h, 24h")
doc.add_paragraph("• Enkripsi Aman: Fernet (AES-128) + PBKDF2-HMAC-SHA256")
doc.add_paragraph("• Configuration Diff: Deteksi perubahan otomatis")
doc.add_paragraph("• Retention Policy: Penyimpanan 365 hari")
doc.add_paragraph("• Windows Service: Berjalan sebagai background service")
doc.add_paragraph("• Portable EXE: Distribusi single executable")

# Add diagram section
doc.add_heading('System Architecture', level=2)
doc.add_paragraph("Diagram arsitektur sistem tersedia di file flowchart/system_architecture.mmd")
doc.add_paragraph("Konversi ke PNG dengan VS Code + ekstensi Mermaid")

doc.add_heading('Backup Process Flow', level=2)
doc.add_paragraph("Diagram alir proses backup tersedia di file flowchart/backup_process.mmd")

# Save document
doc.save('report_without_images.docx')
print(f"Laporan tanpa gambar berhasil dibuat: {os.path.abspath('report_without_images.docx')}")
